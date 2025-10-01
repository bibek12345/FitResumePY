 codex/build-fitresume-project-with-full-features-nhskwq
from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path

from ..config import get_settings
from ..utils.docx_utils import (
    RenderContext,
    create_placeholder_template,
    render_resume_docx,
)
from .rewrite_service import RewriteResult


class ArtifactService:
    TEMPLATE_VERSION = "1.0"

    def __init__(self, artifacts_root: Path | None = None, template_path: Path | None = None) -> None:
        settings = get_settings()
        self.artifacts_root = artifacts_root or settings.artifacts_root
        self.template_path = template_path or Path("templates/resume_template.docx").resolve()
        self._ensure_template_exists()

    def create_artifact(
        self,
        *,
        company_name: str,
        job_key: str,
        rewrite_result: RewriteResult,
    ) -> Path:
        folder = self._resolve_folder(company_name, job_key)
        folder.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now(timezone.utc)
        filename = f"resume_{timestamp.strftime('%Y%m%d_%H%M%S')}.docx"
        artifact_path = folder / filename

        context = self._build_context(rewrite_result)
        render_resume_docx(artifact_path, context)

        meta_path = folder / "meta.json"
        meta_payload = {
            "company": company_name,
            "job_key": job_key,
            "artifact": artifact_path.name,
            "created_at": timestamp.isoformat(),
            "template_version": self.TEMPLATE_VERSION,
            "model_name": rewrite_result.model_name,
            "prompt_hash": rewrite_result.prompt_hash,
            "plan": rewrite_result.plan,
            "mock_output": rewrite_result.mock,
        }
        if rewrite_result.token_usage:
            meta_payload["token_usage"] = rewrite_result.token_usage
        meta_path.write_text(json.dumps(meta_payload, indent=2), encoding="utf-8")
        return artifact_path

    def _resolve_folder(self, company_name: str, job_key: str) -> Path:
        folder_name = f"{self._sanitize(company_name)}__{self._sanitize(job_key)}"
        return self.artifacts_root / folder_name

    def _sanitize(self, value: str) -> str:
        return re.sub(r"[^A-Za-z0-9_-]", "_", value.strip() or "job")[:60]

    def _build_context(self, rewrite_result: RewriteResult) -> RenderContext:
        plan = rewrite_result.plan
        skills = [str(skill) for skill in plan.get("skills", []) if str(skill).strip()]
        experience_blocks: list[str] = []
        for exp in plan.get("experience", []):
            header = f"{exp.get('employer', '')} — {exp.get('role', '')} ({exp.get('start', '')} - {exp.get('end', '')})"
            bullets = [f"• {bullet}" for bullet in exp.get("bullets", []) if str(bullet).strip()]
            block_lines = [header.strip()] if header.strip() else []
            block_lines.extend(bullets)
            if block_lines:
                experience_blocks.append("\n".join(block_lines))
        education = [str(item) for item in plan.get("education", []) if str(item).strip()]
        certifications = [str(item) for item in plan.get("certifications", []) if str(item).strip()]
        return RenderContext(
            summary=str(plan.get("summary", "")),
            skills=skills,
            experience=experience_blocks,
            education=education,
            certifications=certifications,
        )

    def _ensure_template_exists(self) -> None:
        create_placeholder_template(self.template_path)

"""Services for generating resume artifacts from templates."""
from __future__ import annotations

from pathlib import Path
from typing import Any, Mapping

from docxtpl import DocxTemplate
from docx import Document


class ArtifactService:
    """Service responsible for rendering DOCX resume artifacts."""

    def __init__(self, template_path: str | Path | None = None) -> None:
        self.template_path = Path(template_path or "templates/resume_template.docx")
        self._ensure_template_exists()

    def _ensure_template_exists(self) -> None:
        """Create the DOCX template if it is missing from disk.

        The project historically stored a binary template in version control.
        This helper recreates the template at runtime when it is not present so
        artifact generation can continue without shipping binary assets.
        """

        if self.template_path.exists():
            return

        self.template_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph("{{ summary }}")

        document.add_heading("Skills", level=1)
        document.add_paragraph("{{ skills }}")

        document.add_heading("Experience", level=1)
        document.add_paragraph("{{ experience }}")

        document.add_heading("Education", level=1)
        document.add_paragraph("{{ education }}")

        document.add_heading("Certifications", level=1)
        document.add_paragraph("{{ certifications }}")

        document.save(self.template_path)

    def render_artifact(
        self, context: Mapping[str, Any], destination: str | Path
    ) -> Path:
        """Render the resume artifact into ``destination`` using the template.

        Parameters
        ----------
        context:
            A mapping with the keys expected by the template (``summary``,
            ``skills``, ``experience``, ``education``, ``certifications``).
        destination:
            Where the rendered document should be written.
        """

        template = DocxTemplate(str(self.template_path))
        template.render(dict(context))

        destination_path = Path(destination)
        destination_path.parent.mkdir(parents=True, exist_ok=True)
        template.save(str(destination_path))
        return destination_path
 main
